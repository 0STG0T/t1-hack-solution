from typing import Dict, Any, Optional
import docx
import io
from .base_processor import BaseProcessor

class DOCXProcessor(BaseProcessor):
    def __init__(self):
        super().__init__()
        self._preview_length = 500

    @property
    def content_type(self) -> str:
        return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    async def process_content(self, content: bytes, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Process DOCX content and extract text with metadata"""
        try:
            # Create document object from bytes
            doc_stream = io.BytesIO(content)
            doc = docx.Document(doc_stream)

            # Extract all text content with formatting
            full_text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text += para.text.strip() + "\n"

            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        full_text += row_text + "\n"

            # Clean and normalize text
            full_text = self._clean_text(full_text)

            # Extract document properties and styles
            doc_info = {
                'title': metadata.get('title', 'Untitled Document'),
                'paragraph_count': len(doc.paragraphs),
                'section_count': len(doc.sections),
                'table_count': len(doc.tables),
                'has_headers': any(section.header.paragraphs[0].text.strip()
                                 for section in doc.sections),
                'has_footers': any(section.footer.paragraphs[0].text.strip()
                                 for section in doc.sections),
                'styles_used': self._extract_styles(doc)
            }

            # Merge with provided metadata
            doc_metadata = {
                **metadata,
                **doc_info,
                'source_type': 'docx',
                'processor_version': docx.__version__
            }

            return {
                'content': full_text,
                'metadata': doc_metadata
            }

        except Exception as e:
            raise ValueError(f"Error processing DOCX content: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        # Remove excessive whitespace and normalize line endings
        text = ' '.join(text.split())
        # Remove any control characters
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
        return text.strip()

    def _extract_styles(self, doc) -> Dict[str, Any]:
        """Extract document style information"""
        styles = {}
        try:
            for style in doc.styles:
                if style.type == docx.enum.style.WD_STYLE_TYPE.PARAGRAPH:
                    styles[style.name] = {
                        'type': 'paragraph',
                        'is_custom': style.custom
                    }
        except:
            pass
        return styles
