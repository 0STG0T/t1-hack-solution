import io
import fitz  # PyMuPDF
from docx import Document
from typing import Dict, Any, Optional, Tuple
from .base_processor import BaseProcessor
import chardet
from datetime import datetime

class PDFProcessor(BaseProcessor):
    @property
    def content_type(self) -> str:
        return 'application/pdf'

    async def process_content(self, content: bytes) -> Dict[str, Any]:
        try:
            pdf_document = fitz.open(stream=content, filetype="pdf")
            metadata = self._extract_metadata(pdf_document)
            text_content, structure = self._extract_content(pdf_document)

            return {
                'title': metadata.get('title', 'Untitled PDF Document'),
                'text': text_content,
                'metadata': metadata,
                'structure': structure,
                'success': True
            }
        except Exception as e:
            return {
                'error': f"PDF processing failed: {str(e)}",
                'success': False
            }
        finally:
            if 'pdf_document' in locals():
                pdf_document.close()

    def _extract_metadata(self, doc) -> Dict[str, Any]:
        metadata = doc.metadata
        return {
            'title': metadata.get('title', ''),
            'author': metadata.get('author', ''),
            'subject': metadata.get('subject', ''),
            'keywords': metadata.get('keywords', ''),
            'creator': metadata.get('creator', ''),
            'producer': metadata.get('producer', ''),
            'creation_date': metadata.get('creationDate', ''),
            'modification_date': metadata.get('modDate', ''),
            'page_count': doc.page_count
        }

    def _extract_content(self, doc) -> Tuple[str, Dict[str, Any]]:
        text_content = []
        structure = {'headings': [], 'paragraphs': 0, 'images': 0}

        for page_num, page in enumerate(doc):
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("lines"):
                    text = ' '.join(span["text"] for line in block["lines"]
                                  for span in line["spans"])
                    text_content.append(text)

                    # Detect headings by font size
                    if any(span["size"] > 12 for line in block["lines"]
                          for span in line["spans"]):
                        structure['headings'].append({
                            'text': text,
                            'page': page_num + 1
                        })
                    structure['paragraphs'] += 1
                elif block.get("image", 0):
                    structure['images'] += 1

        return ' '.join(text_content), structure

class DOCXProcessor(BaseProcessor):
    @property
    def content_type(self) -> str:
        return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    async def process_content(self, content: bytes) -> Dict[str, Any]:
        try:
            doc = Document(io.BytesIO(content))
            metadata = self._extract_metadata(doc)
            text_content, structure = self._extract_content(doc)

            return {
                'title': metadata.get('title', 'Untitled DOCX Document'),
                'text': text_content,
                'metadata': metadata,
                'structure': structure,
                'success': True
            }
        except Exception as e:
            return {
                'error': f"DOCX processing failed: {str(e)}",
                'success': False
            }

    def _extract_metadata(self, doc) -> Dict[str, Any]:
        core_props = doc.core_properties
        return {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'subject': core_props.subject or '',
            'keywords': core_props.keywords or '',
            'created': str(core_props.created) if core_props.created else '',
            'modified': str(core_props.modified) if core_props.modified else '',
            'category': core_props.category or '',
            'comments': core_props.comments or ''
        }

    def _extract_content(self, doc) -> Tuple[str, Dict[str, Any]]:
        text_content = []
        structure = {'headings': [], 'paragraphs': 0, 'tables': 0}

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_content.append(text)
                structure['paragraphs'] += 1
                if para.style.name.startswith('Heading'):
                    structure['headings'].append({
                        'level': int(para.style.name[-1]),
                        'text': text
                    })

        structure['tables'] = len(doc.tables)
        return ' '.join(text_content), structure

class TXTProcessor(BaseProcessor):
    @property
    def content_type(self) -> str:
        return 'text/plain'

    async def process_content(self, content: bytes) -> Dict[str, Any]:
        try:
            # Detect encoding
            encoding = chardet.detect(content)['encoding'] or 'utf-8'
            text = content.decode(encoding)

            metadata = self._extract_metadata(content, encoding)
            text_content, structure = self._extract_content(text)

            return {
                'title': structure['title'],
                'text': text_content,
                'metadata': metadata,
                'structure': structure,
                'success': True
            }
        except Exception as e:
            return {
                'error': f"TXT processing failed: {str(e)}",
                'success': False
            }

    def _extract_metadata(self, content: bytes, encoding: str) -> Dict[str, Any]:
        return {
            'size': len(content),
            'encoding': encoding,
            'timestamp': datetime.now().isoformat(),
            'line_count': content.count(b'\n') + 1
        }

    def _extract_content(self, text: str) -> Tuple[str, Dict[str, Any]]:
        lines = text.split('\n')
        structure = {
            'title': next((line.strip() for line in lines if line.strip()), 'Untitled TXT Document'),
            'paragraphs': sum(1 for line in lines if line.strip()),
            'empty_lines': sum(1 for line in lines if not line.strip())
        }
        return text, structure

def get_processor_for_content_type(content_type: str, mongodb_client) -> BaseProcessor:
    """Factory function to get the appropriate processor for a content type"""
    processors = {
        'application/pdf': PDFProcessor,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DOCXProcessor,
        'text/plain': TXTProcessor
    }

    processor_class = processors.get(content_type)
    if not processor_class:
        raise ValueError(f"Unsupported content type: {content_type}")

    return processor_class(mongodb_client)
